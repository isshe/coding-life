# coding=utf-8

import sys

import pymysql
from docx import Document

def get_db_table_list(cursor):
    table_list = []
    sql = 'show tables'
    cursor.execute(sql)
    infos = cursor.fetchall()
    for info in infos:
        table_list.append(info[0])
    return table_list


def get_sparta_waf_conn(host):
    conn = pymysql.connect(
        host=host,
        port=3306,
        user='root',
        password='chudai',
        db='sparta_waf',
        charset="utf8")
    return conn


def get_waf_host_config_conn(host):
    conn = pymysql.connect(
        host=host,
        port=3306,
        user='root',
        password='chudai',
        db='waf_host_config',
        charset="utf8")
    return conn


def get_table_desc_list(cursor, table):
    sql = "SHOW FULL COLUMNS FROM " + table
    cursor.execute(sql)
    infos = cursor.fetchall()
    desc_list = []
    for info in infos:
        field = {}
        field['name'] = info[0]
        field['type'] = info[1]
        field['null'] = info[3]
        field['default'] = info[5]
        field['comment'] = info[8]
        # print(info[8])
        desc_list.append(field)
    return desc_list

def get_db_table_infos(cursor):
    spt_table_list = get_db_table_list(cursor)

    table_infos = {}
    for table in spt_table_list:
        table_infos[table] = get_table_desc_list(cursor, table)
    return table_infos
    # print(table_infos)


def to_docx(infos, header, filename):
    doc = Document()
    doc.add_heading(header,0)
    for table_name, info_array in infos.items():
        doc.add_heading(table_name, level=2)
        #tlen = len(info_array) + 1
        table = doc.add_table(rows=1, cols=5, style='Table Grid')
        cells = table.rows[0].cells
        # 为每一个单元格赋值
        # 注：值都要为字符串类型
        cells[0].text = u'名称'
        cells[1].text = u'类型'
        cells[2].text = u'说明'
        cells[3].text = u'默认值'
        cells[4].text = u'允许空(是/否)'

        for info in info_array:
            # 为表格添加一行
            new_cells = table.add_row().cells
            new_cells[0].text = info['name']
            new_cells[1].text = info['type']
            new_cells[2].text = info['comment']
            new_cells[3].text = info['default'] or ''
            new_cells[4].text = u"是"
            if info['null'] != "YES" :
                new_cells[4].text = u"否"
        doc.add_page_break()
    doc.save(filename)



if __name__ == "__main__":
    reload(sys)
    sys.setdefaultencoding('utf-8')

    host = '10.16.71.44'

    sptconn = get_sparta_waf_conn(host)
    sptcursor = sptconn.cursor()
    sptinfos = get_db_table_infos(sptcursor)
    # print(sptinfos)
    to_docx(sptinfos, "sparta_waf", './sparta_waf.docx')

    whcconn = get_waf_host_config_conn(host)
    whccursor = whcconn.cursor()
    whcinfos = get_db_table_infos(whccursor)
    to_docx(whcinfos, "waf_host_config", './waf_host_config.docx')

